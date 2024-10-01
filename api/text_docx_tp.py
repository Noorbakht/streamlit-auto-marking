from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

import random

random_id = f'Tender No. : TO{str(random.randint(2000000, 9999999))}'

def add_header(doc):
    # Add a header section like the PDF
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    run = header_paragraph.add_run()
    
    # Add the logo image on the left
    run.add_picture('tp_icon.png', width=Inches(1.8))  # Adjust size as needed

    # Add space before the header text
    header_paragraph.add_run("                                                                                    ")

    # Add the header text on the right
    run = header_paragraph.add_run(random_id)
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add blank space within the header
    header_paragraph.add_run("\n")

def add_footer(doc):
    
    # Add a footer section
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]

    # Add blank space within the footer
    run = footer_paragraph.add_run()
    footer_paragraph.add_run("\n")
    
    # footer_paragraph.text = "SECTION 3A: STUDENT INFORMATION SYSTEM REQUIREMENTS"
    # footer_paragraph.style.font.name = 'Arial'
    # footer_paragraph.style.font.size = Pt(10)
    # footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def cm_to_pt(cm):
    return cm * 28.3464566929134

def gen_docx(html_text, line_height=12, font_size=11):
    doc = Document()

    # Set margins for the document (similar to PDF margins)
    section = doc.sections[0]
    section.top_margin = Pt(5)
    section.left_margin = Pt(60)
    section.right_margin = Pt(60)
    section.bottom_margin = Pt(5)
    
    # Add Header
    add_header(doc)

    # Add Footer
    add_footer(doc)
    
    # Data preparation for DOCX output
    TABLE_DATA = []
    rows = html_text.split("\n")

    for row in rows[0:-1]:
        cells = row.split("|")
        number = cells[1].strip()
        spec = cells[2].strip()
        TABLE_DATA.append((number, spec))

    # Add the main content of the document as a multi-level list
    first_section = True

    for d in range(len(TABLE_DATA)):
        row = TABLE_DATA[d]
        
        # Section titles (Numbered list, Level 1)
        if len(row[0]) == 1:
            if not first_section:
                doc.add_page_break()
            else:
                first_section = False
            p = doc.add_paragraph(f"{row[1]}.")  # Use built-in ListNumber style
            p.paragraph_format.space_before = Pt(10)
            # p.paragraph_format.space_after = Pt(line_height)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.style.paragraph_format.level = 0
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(font_size)
                run.bold = True  # Make the text bold
        
        # Sub-points (Numbered list, Level 2)
        elif row[0][0] != '-':
            p = doc.add_paragraph(f"{row[0]}. {row[1]}")  # Use built-in ListNumber style
            p.paragraph_format.space_after = Pt(line_height)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # Set level for sub-points
            p.style.paragraph_format.level = 1
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(font_size)

    return doc