from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_header(doc):
    # Add a header section like the PDF
    header = doc.sections[0].header
    header_paragraph = header.paragraphs[0]
    
    # Center align and set font for header
    run = header_paragraph.add_run("Official (Closed), Non-Sensitive")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add blank space within the header
    header_paragraph.add_run("\n\n")

def add_footer(doc):
    # Add a footer section
    footer = doc.sections[0].footer
    footer_paragraph = footer.paragraphs[0]
    
    # Add "Singapore Polytechnic" on the bottom left
    footer_paragraph.text = "Singapore Polytechnic"
    footer_paragraph.style.font.name = 'Times New Roman'
    footer_paragraph.style.font.size = Pt(10)
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def cm_to_pt(cm):
    return cm * 28.3464566929134

def gen_docx(html_text, line_height=12, font_size=12):
    doc = Document()

    # Set margins for the document (similar to PDF margins)
    section = doc.sections[0]
    section.top_margin = Pt(cm_to_pt(2.54))
    section.left_margin = Pt(cm_to_pt(3.17))
    section.right_margin = Pt(cm_to_pt(3.17))
    section.bottom_margin = Pt(cm_to_pt(2.54))
    
    # Add Header
    add_header(doc)

    # Add Footer
    add_footer(doc)
    
    # Data preparation for DOCX output
    TABLE_DATA = []
    rows = html_text.split("\n")

    for row in rows[0:-1]:
        cells = row.split("|")
        number = cells[1].strip()
        spec = cells[2].strip()
        TABLE_DATA.append((number, spec))

    # Add the main content of the document as a multi-level list
    for d in range(len(TABLE_DATA)):
        row = TABLE_DATA[d]
        
        # Section titles (Numbered list, Level 1)
        if len(row[0]) == 1:
            p = doc.add_paragraph(f"{row[1]}")  # Use built-in ListNumber style
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(line_height)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            p.style.paragraph_format.level = 0
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(font_size)
        
        # Sub-points (Numbered list, Level 2)
        elif row[0][0] != '-':
            p = doc.add_paragraph(f"{row[0]} {row[1]}")  # Use built-in ListNumber style
            p.paragraph_format.left_indent = Pt(20)  # Indent for Level 2
            p.paragraph_format.space_after = Pt(line_height)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # Set level for sub-points
            p.style.paragraph_format.level = 1
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(font_size)

    return doc